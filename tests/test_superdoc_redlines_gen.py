"""yuch85/superdoc-redlines DOCX redline generator (SuperDoc headless CLI)."""

from __future__ import annotations

import shutil
import subprocess
import zipfile
from pathlib import Path

import pytest
from helpers import CORPUS

from neurotic_docx_bench import superdoc_redlines_gen

MANIFEST = CORPUS / "centralized_mapping.csv"
SOURCE = CORPUS / "docx_source"
REPO = Path(__file__).resolve().parents[1] / "superdoc-redlines"

_HAVE_TOOL = (REPO / "superdoc-redline.mjs").is_file() and (REPO / "node_modules").is_dir()
_HAVE_NODE = shutil.which("node") is not None

def _cli_runtime_broken() -> str | None:
    """Probe the vendor CLI once; return why it is unusable, or None if it works.

    Plan Chapter 6 D6: this machine runs a bleeding-edge Node, and the vendor's
    transitive dependency (@harbour-enterprises/superdoc) dies there with
    ``TypeError: varStorage.getItem is not a function``. Skipping with the reason
    NAMED keeps the situation visible; letting the suite sit permanently red
    hides real regressions behind a known one, and silently passing would hide it
    entirely. This is a runtime-compatibility gap, and whether it counts against
    the vendor is decided on a supported Node — not here.
    """
    if not (_HAVE_TOOL and _HAVE_NODE):
        return "superdoc-redlines clone (repo root, npm-installed) or node absent"
    probe = subprocess.run(
        ["node", str(REPO / "superdoc-redline.mjs"), "--help"],
        capture_output=True, text=True, cwd=str(REPO), check=False,
    )
    blob = (probe.stderr or "") + (probe.stdout or "")
    if "varStorage.getItem is not a function" in blob:
        node_v = subprocess.run(
            ["node", "--version"], capture_output=True, text=True, check=False,
        ).stdout.strip()
        return (
            f"superdoc-redlines CLI is unusable on this runtime ({node_v}): its dependency "
            f"@harbour-enterprises/superdoc raises 'varStorage.getItem is not a function'. "
            f"Vendor declares engines.node >=18.0.0. See plan Chapter 6 D6 — benchmark it "
            f"on a supported Node LTS before attributing this to the vendor."
        )
    return None


_CLI_BROKEN = _cli_runtime_broken()

requires_tool = pytest.mark.skipif(_CLI_BROKEN is not None, reason=_CLI_BROKEN or "")
requires_manifest = pytest.mark.skipif(not MANIFEST.is_file(), reason="corpus manifest absent")


def test_build_edits_aligns_blocks():
    """Pure alignment: replace/delete/insert ops from block-text sequences."""
    base = [
        {"seqId": "b001", "text": "Title", "type": "paragraph"},
        {"seqId": "b002", "text": "Old body.", "type": "paragraph"},
        {"seqId": "b003", "text": "Removed paragraph.", "type": "paragraph"},
        {"seqId": "b004", "text": "Common tail.", "type": "paragraph"},
    ]
    nxt = [
        {"seqId": "b001", "text": "Title", "type": "paragraph"},
        {"seqId": "b002", "text": "New body.", "type": "paragraph"},
        {"seqId": "b003", "text": "Common tail.", "type": "paragraph"},
        {"seqId": "b004", "text": "Appended paragraph.", "type": "paragraph"},
    ]
    payload = superdoc_redlines_gen.build_edits(base, nxt, author="t")
    ops = [(e["operation"], e.get("blockId") or e.get("afterBlockId")) for e in payload["edits"]]
    assert ("replace", "b002") in ops
    assert ("delete", "b003") in ops
    assert any(op == "insert" for op, _ in ops)


def test_build_edits_empty_base_raises_cleanly():
    """A base doc with no addressable blocks must fail with a clear message,
    not an IndexError."""
    nxt = [{"seqId": "b001", "text": "New paragraph.", "type": "paragraph"}]
    with pytest.raises(ValueError, match="no addressable blocks"):
        superdoc_redlines_gen.build_edits([], nxt, author="t")


@requires_tool
def test_generate_one_emits_tracked_changes(tmp_path):
    """Synthetic base/next → DOCX must carry w:ins and/or w:del."""
    from docx import Document

    base = tmp_path / "base.docx"
    nxt = tmp_path / "next.docx"
    out = tmp_path / "out.docx"

    b = Document()
    b.add_paragraph("The quick brown fox jumps over the lazy dog.")
    b.save(str(base))
    n = Document()
    n.add_paragraph("The quick brown fox walks past the lazy dog.")
    n.save(str(nxt))

    superdoc_redlines_gen.generate_one(
        base, nxt, out, repo=REPO, author="superdoc-redlines", workdir=tmp_path / "work",
    )
    assert out.is_file()
    with zipfile.ZipFile(out) as z:
        xml = z.read("word/document.xml").decode("utf-8", "ignore")
    assert "<w:ins" in xml or "<w:del" in xml


@requires_tool
@requires_manifest
def test_run_batch_produces_tracked_redline(tmp_path):
    ok, failed, _timings = superdoc_redlines_gen.run_batch(
        out=tmp_path / "docx",
        manifest=MANIFEST,
        source_dir=SOURCE,
        statuses={"ok"},
        limit=1,
        tool="superdoc-redlines",
        author="superdoc-redlines",
        force=True,
        repo=REPO,
        jobs=2,
    )
    assert ok >= 1, failed
    outs = list((tmp_path / "docx").glob("*_superdoc-redlines_redline.docx"))
    assert len(outs) == 1
    with zipfile.ZipFile(outs[0]) as z:
        xml = z.read("word/document.xml").decode("utf-8", "ignore")
    assert "w:document" in xml
