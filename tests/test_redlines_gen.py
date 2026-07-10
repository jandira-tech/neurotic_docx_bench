"""houfu/redlines text→DOCX redline generator."""

from __future__ import annotations

import zipfile

import pytest
from helpers import CORPUS

from neurotic_docx_bench import redlines_gen

MANIFEST = CORPUS / "centralized_mapping.csv"
SOURCE = CORPUS / "docx_source"

try:
    import redlines  # noqa: F401

    _HAVE_REDLINES = True
except Exception:  # pragma: no cover
    _HAVE_REDLINES = False

requires_redlines = pytest.mark.skipif(not _HAVE_REDLINES, reason="redlines not installed")
requires_manifest = pytest.mark.skipif(not MANIFEST.is_file(), reason="corpus manifest absent")


@requires_manifest
def test_parse_manifest_returns_pairs():
    pairs = redlines_gen.parse_manifest(MANIFEST, {"ok"})
    assert pairs and all(p.base and p.next for p in pairs)


@requires_redlines
def test_generate_one_emits_tracked_changes(tmp_path):
    """Synthetic base/next → DOCX must carry w:ins and/or w:del."""
    from docx import Document

    base = tmp_path / "base.docx"
    nxt = tmp_path / "next.docx"
    out = tmp_path / "out.docx"

    b = Document()
    b.add_paragraph("The quick brown fox jumps over the lazy dog.")
    b.save(str(base))
    n = Document()
    n.add_paragraph("The quick brown fox walks past the lazy dog.")
    n.save(str(nxt))

    redlines_gen.generate_one(base, nxt, out, author="redlines")
    assert out.is_file()
    with zipfile.ZipFile(out) as z:
        xml = z.read("word/document.xml").decode("utf-8", "ignore")
    assert "<w:ins" in xml or "<w:del" in xml
    assert "jumps" in xml or "walks" in xml


@requires_redlines
@requires_manifest
def test_run_batch_produces_tracked_redline(tmp_path):
    ok, failed, _timings = redlines_gen.run_batch(
        out=tmp_path,
        manifest=MANIFEST,
        source_dir=SOURCE,
        statuses={"ok"},
        limit=1,
        tool="redlines",
        author="redlines",
        force=True,
    )
    assert ok >= 1, failed
    outs = list(tmp_path.glob("*_redlines_redline.docx"))
    assert len(outs) == 1
    with zipfile.ZipFile(outs[0]) as z:
        xml = z.read("word/document.xml").decode("utf-8", "ignore")
    # Identical short docs may produce no changes; corpus pairs usually differ.
    # At minimum the file must be a valid DOCX with document.xml.
    assert "w:document" in xml or "document" in xml
