"""Generate tracked-change redlines via the houfu/redlines text library.

``redlines`` compares plain text (not OOXML). This adapter:

1. Extracts paragraph text from base/next DOCX (python-docx).
2. Runs ``Redlines`` (preferring ``NupunktProcessor`` when installed — better for
   legal abbreviations / citations).
3. Writes a new DOCX whose body is the change stream as Word ``w:ins`` / ``w:del``
   markup so the soffice script_redlines pipeline can score it against the Word oracle.

This is an intentional *text-level* baseline: formatting, tables, images, and
structure are not preserved. Scores measure how far a pure text redliner gets
toward Word's document-level tracked changes.

Usage:
  uv run python -m neurotic_docx_bench.redlines_gen --out $RUN_DIR/docx --tool redlines \
    [--manifest corpus/word_based/centralized_mapping.csv] \
    [--source-dir corpus/word_based/docx_source] [--limit N]
"""

from __future__ import annotations

import argparse
import csv
import json
import time
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from redlines import Redlines
from redlines.processor import WholeDocumentProcessor

try:
    from redlines.processor import NupunktProcessor

    _PROCESSOR = NupunktProcessor()
except Exception:  # pragma: no cover — nupunkt optional at import time
    _PROCESSOR = WholeDocumentProcessor()


@dataclass
class Pair:
    base: str
    next: str


def output_name(pair: Pair, tool: str) -> str:
    """Canonical candidate name: ``<base>_<next>_<tool>_redline.docx``."""
    return f"{pair.base}_{pair.next}_{tool}_redline.docx"


def parse_manifest(csv_path: Path, statuses: set[str]) -> list[Pair]:
    """Return pairs whose batch_status is wanted (empty status = include)."""
    pairs: list[Pair] = []
    with Path(csv_path).open(newline="") as fh:
        for row in csv.DictReader(fh):
            base = (row.get("base") or "").strip()
            nxt = (row.get("next") or "").strip()
            status = (row.get("batch_status") or "").strip()
            if not base or not nxt:
                continue
            if statuses and status and status not in statuses:
                continue
            pairs.append(Pair(base=base, next=nxt))
    return pairs


def extract_text(docx_path: Path) -> str:
    """Plain-text extraction: paragraph bodies joined by newlines.

    Tables / headers / footers are omitted — redlines is a string differ; the
    adapter stays honest about that scope.
    """
    doc = Document(str(docx_path))
    return "\n".join(p.text for p in doc.paragraphs)


def _make_run(text: str, *, delete: bool) -> OxmlElement:
    r = OxmlElement("w:r")
    if delete:
        t = OxmlElement("w:delText")
    else:
        t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _append_revision(
    paragraph,
    text: str,
    *,
    kind: str,
    rev_id: int,
    author: str,
    when: str,
) -> None:
    """Append equal run, ``w:ins``, or ``w:del`` for a non-empty text slice."""
    if not text:
        return
    if kind == "equal":
        paragraph.add_run(text)
        return
    tag = "w:ins" if kind == "insert" else "w:del"
    el = OxmlElement(tag)
    el.set(qn("w:id"), str(rev_id))
    el.set(qn("w:author"), author)
    el.set(qn("w:date"), when)
    el.append(_make_run(text, delete=(kind == "delete")))
    paragraph._p.append(el)


def _emit_change_stream(
    doc: Document,
    changes: list[dict],
    *,
    author: str,
) -> None:
    """Render structured redlines changes into ``doc`` as tracked runs.

    Newlines in change text become paragraph breaks so multi-paragraph sources
    stay roughly paragraph-aligned.
    """
    when = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
    # python-docx may start with zero paragraphs depending on version/template —
    # always create the first body paragraph ourselves.
    para = doc.add_paragraph()
    rev_id = 0

    def new_paragraph() -> None:
        nonlocal para
        para = doc.add_paragraph()

    def emit(text: str, kind: str) -> None:
        nonlocal rev_id
        if text is None:
            return
        parts = text.split("\n")
        for i, part in enumerate(parts):
            if i > 0:
                new_paragraph()
            if not part:
                continue
            if kind in ("insert", "delete"):
                rev_id += 1
                _append_revision(
                    para, part, kind=kind, rev_id=rev_id, author=author, when=when,
                )
            else:
                _append_revision(
                    para, part, kind="equal", rev_id=rev_id, author=author, when=when,
                )

    for ch in changes:
        ctype = ch.get("type")
        if ctype == "equal":
            emit(ch.get("text") or "", "equal")
        elif ctype == "insert":
            emit(ch.get("text") or "", "insert")
        elif ctype == "delete":
            emit(ch.get("text") or "", "delete")
        elif ctype == "replace":
            emit(ch.get("source_text") or "", "delete")
            emit(ch.get("test_text") or "", "insert")
        else:
            # Unknown op — best-effort surface whatever text fields exist.
            emit(ch.get("text") or ch.get("source_text") or "", "equal")
            if ch.get("test_text"):
                emit(ch["test_text"], "insert")


def generate_one(base_path: Path, next_path: Path, out_path: Path, *, author: str) -> None:
    """base.docx + next.docx → tracked-change redline DOCX via redlines text diff."""
    source = extract_text(base_path)
    test = extract_text(next_path)
    rl = Redlines(source, test, processor=_PROCESSOR)
    # output_json is the structured contract; parse for change ops.
    payload = json.loads(rl.output_json())
    changes = payload.get("changes") or []

    doc = Document()
    _emit_change_stream(doc, changes, author=author)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def run_batch(
    *,
    out: Path,
    manifest: Path,
    source_dir: Path,
    statuses: set[str],
    limit: int | None,
    tool: str,
    author: str,
    force: bool,
) -> tuple[int, list[dict], dict[str, int]]:
    out.mkdir(parents=True, exist_ok=True)
    pairs = parse_manifest(manifest, statuses)
    if limit:
        pairs = pairs[:limit]
    ok = 0
    failed: list[dict] = []
    timings: dict[str, int] = {}
    for pair in pairs:
        doc = f"{pair.base}_{pair.next}"
        name = output_name(pair, tool)
        out_path = out / name
        if not force and out_path.exists():
            ok += 1
            continue
        base_path = source_dir / f"{pair.base}.docx"
        next_path = source_dir / f"{pair.next}.docx"
        if not base_path.exists() or not next_path.exists():
            failed.append({"doc": doc, "stage": "missing_source", "error": "source docx not found"})
            continue
        try:
            t0 = time.perf_counter_ns()
            generate_one(base_path, next_path, out_path, author=author)
            timings[name.replace(".docx", "")] = time.perf_counter_ns() - t0
            ok += 1
        except Exception as exc:  # one bad pair must not stop the batch
            failed.append({"doc": doc, "stage": "generate", "error": str(exc)})
    return ok, failed, timings


def main(argv: list[str] | None = None) -> int:
    import os

    p = argparse.ArgumentParser(description="redlines (houfu) text→DOCX redline generator")
    default_out = (
        os.path.join(os.environ["RUN_DIR"], "docx") if os.environ.get("RUN_DIR") else "out/docx"
    )
    p.add_argument("--out", default=default_out)
    p.add_argument("--manifest", default="corpus/word_based/centralized_mapping.csv")
    p.add_argument("--source-dir", default="corpus/word_based/docx_source")
    p.add_argument("--status", default="ok")
    p.add_argument("--limit", type=int, default=None)
    p.add_argument("--tool", default="redlines")
    p.add_argument("--author", default="redlines")
    p.add_argument("--force", action="store_true")
    args = p.parse_args(argv)

    ok, failed, timings = run_batch(
        out=Path(args.out),
        manifest=Path(args.manifest),
        source_dir=Path(args.source_dir),
        statuses=set(args.status.split(",")) if args.status else set(),
        limit=args.limit,
        tool=args.tool,
        author=args.author,
        force=args.force,
    )
    out_dir = Path(args.out)
    (out_dir.parent / "generate_failures.json").write_text(json.dumps(failed, indent=2))
    (out_dir.parent / "generate_timings.json").write_text(json.dumps(timings))
    print(f"[redlines] wrote {ok} redline(s) → {args.out}")
    if failed:
        print(f"[redlines] {len(failed)} pair(s) skipped:")
        for f in failed[:10]:
            print(f"  {f['doc']} [{f['stage']}]: {f['error']}")
    return 0 if ok > 0 else 1


if __name__ == "__main__":
    raise SystemExit(main())
